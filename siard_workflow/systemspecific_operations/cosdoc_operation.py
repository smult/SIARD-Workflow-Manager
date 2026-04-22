"""
siard_workflow/systemspecific_operations/cosdoc_operation.py

CosDocMailMergeOperation — avkryptering og mailmerge for CosDoc SIARD-arkiver.

CosDoc lagrer dokumenter i tabellen Eef_ElFiler:
  - Eef_ElFil   (BLOB)  : selve filen, passordbeskyttet
  - Eef_FilNavn (NCHAR) : filnavn, brukes til å utlede passord
  - Eef_SeqNr   (INT)   : 1 = hoveddokument, 2 = flettefil (datakilde)
  - Eef_EveID   (INT)   : kobling mellom hoveddokument og flettefil

Passord utledes fra Eef_FilNavn:
  filstamme (uten endelse) → snu rekkefølge → legg R foran
  Eks: "201312345.doc" → "543213102" → "R543213102"

Prosess per dokumentpar (samme Eef_EveID):
  1. Trekk ut begge blob-filer fra SIARD lob-mappen
  2. Avkrypter med utledet passord (msoffcrypto-tool)
  3. Konverter begge til DOCX (LibreOffice)
  4. Les merge-data fra datakilde-DOCX (første tabell i dokumentet)
  5. Flett hoveddokument med merge-data (docx-mailmerge2)
  6. Skriv flettet DOCX tilbake til lob-mappen (erstatter gammelt blob)
  7. Slett datakilde-blob og fjern datakilde-rad fra tableX.xml
  8. Oppdater hoveddokument-rad med nytt filnavn, størrelse og digest

Avhengigheter (valgfrie — faller tilbake ved manglende installasjon):
  msoffcrypto-tool : pip install msoffcrypto-tool
  docx-mailmerge2  : pip install docx-mailmerge2
  python-docx      : pip install python-docx  (installeres av mailmerge2)
"""
from __future__ import annotations

import concurrent.futures
import hashlib
import io
import os
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path, PurePosixPath

from siard_workflow.core.base_operation import BaseOperation, OperationResult
from siard_workflow.core.context import WorkflowContext
from siard_workflow.core.siard_format import (
    detect_siard_version, siard_version_transform,
    get_target_siard_version, is_siard_xml,
)


# ── Filtype-klassifisering ───────────────────────────────────────────────────

# Filendelser som kan krypteres med CosDoc og konverteres til DOCX via LO
_WORD_EXTS = frozenset({
    ".doc", ".docx", ".dot", ".dotx",
    ".rtf", ".odt", ".ott",
    ".wpd", ".wp", ".wp5", ".wp6", ".wps", ".lwp",
})


def _is_word_doc(fname: str) -> bool:
    """Returner True hvis filen er et Word-kompatibelt dokument."""
    return Path(fname).suffix.lower() in _WORD_EXTS


# ── Passordderivasjon ─────────────────────────────────────────────────────────

def _derive_password(filename: str) -> str:
    """
    Utled CosDoc-passord fra filnavn.
    Regelen: R + omvendt filstamme (uten filendelse og uten _D-suffiks).

    Flettefiler har "_D" på slutten av filstammen, men passord er basert
    på hoveddokumentets navn — dvs. _D fjernes før reversering.

    Eks: "2013001733.doc"   → stem "2013001733"   → "R3371003102"
         "2013001733_D.doc" → stem "2013001733_D"
                            → strip "_D" → "2013001733" → "R3371003102"
    """
    stem = Path(filename).stem
    if stem.upper().endswith("_D"):
        stem = stem[:-2]
    return "R" + stem[::-1]


# ── MD5-digest ────────────────────────────────────────────────────────────────

def _md5_upper(data: bytes) -> str:
    return hashlib.md5(data).hexdigest().upper()


# ── XML-navnehjelper ──────────────────────────────────────────────────────────

def _local(tag: str) -> str:
    return tag.split("}")[-1] if "}" in tag else tag


def _child_text(elem: ET.Element, tag: str) -> str:
    for child in elem:
        if _local(child.tag) == tag:
            return (child.text or "").strip()
    return ""


def _iter_children(elem: ET.Element, tag: str):
    for child in elem:
        if _local(child.tag) == tag:
            yield child


# ── Metadata-parsing ──────────────────────────────────────────────────────────

def _find_table_info(metadata_bytes: bytes, table_name: str) -> dict | None:
    """
    Finn tabell med gitt navn i metadata.xml.

    Returnerer dict med nøklene:
      schema_folder, table_folder, col_map  ({col_name: col_index})
      lob_folders   ({col_index: lob_folder})  — kun for BLOB-kolonner
    Returnerer None hvis tabellen ikke finnes.
    """
    try:
        root = ET.fromstring(metadata_bytes)
    except ET.ParseError:
        return None

    for schemas_el in root.iter():
        if _local(schemas_el.tag) != "schemas":
            continue
        for schema_el in _iter_children(schemas_el, "schema"):
            schema_folder = _child_text(schema_el, "folder")
            for tables_el in _iter_children(schema_el, "tables"):
                for table_el in _iter_children(tables_el, "table"):
                    name = _child_text(table_el, "name")
                    if name.lower() != table_name.lower():
                        continue
                    table_folder = _child_text(table_el, "folder")
                    col_map: dict[str, int] = {}
                    lob_folders: dict[int, str] = {}
                    col_idx = 0
                    for cols_el in _iter_children(table_el, "columns"):
                        for col_el in _iter_children(cols_el, "column"):
                            col_idx += 1
                            col_name   = _child_text(col_el, "name")
                            lob_folder = _child_text(col_el, "lobFolder")
                            col_map[col_name] = col_idx
                            if lob_folder:
                                lob_folders[col_idx] = lob_folder
                    return {
                        "schema_folder": schema_folder,
                        "table_folder":  table_folder,
                        "col_map":       col_map,
                        "lob_folders":   lob_folders,
                    }
    return None


# ── Tabell-XML-parsing ────────────────────────────────────────────────────────

def _parse_table_rows(xml_bytes: bytes, col_map: dict[str, int]) -> list[dict]:
    """
    Parse tableX.xml og returner liste av rad-dicts.

    Hver dict har:
      _elem      : ET.Element for raden (for videre manipulasjon)
      col_{N}    : tekst-innhold for kolonne N
      blob_{N}   : dict(file, length, digest) for BLOB-kolonner (file-attributt finnes)
    """
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return []

    rows = []
    for row_el in root:
        if _local(row_el.tag) != "row":
            continue
        row: dict = {"_elem": row_el}
        for child in row_el:
            tag = _local(child.tag)
            m = re.match(r"^c(\d+)$", tag, re.IGNORECASE)
            if not m:
                continue
            idx = int(m.group(1))
            if child.get("file"):
                row[f"blob_{idx}"] = {
                    "file":       child.get("file", ""),
                    "length":     child.get("length", "0"),
                    "digestType": child.get("digestType", "MD5"),
                    "digest":     child.get("digest", ""),
                }
            else:
                row[f"col_{idx}"] = (child.text or "").strip()
        rows.append(row)
    return rows


def _get_col(row: dict, idx: int) -> str:
    return row.get(f"col_{idx}", "")


def _get_blob(row: dict, idx: int) -> dict | None:
    return row.get(f"blob_{idx}")


# ── LibreOffice ───────────────────────────────────────────────────────────────

def _find_libreoffice(hint: str = "soffice") -> str | None:
    if hint and shutil.which(hint):
        return hint
    if hint and os.path.isfile(hint):
        return hint
    # Windows søkestier
    if sys.platform == "win32":
        candidates = []
        base = os.environ.get("PROGRAMFILES", r"C:\Program Files")
        if base:
            for sub in ("LibreOffice", "LibreOffice 7", "LibreOffice 24", "OpenOffice"):
                candidates.append(os.path.join(base, sub, "program", "soffice.exe"))
        for base in (r"C:\Program Files", r"C:\Program Files (x86)"):
            if os.path.isdir(base):
                try:
                    for entry in os.listdir(base):
                        if "libre" in entry.lower() or "openoffice" in entry.lower():
                            candidates.append(os.path.join(base, entry, "program", "soffice.exe"))
                except OSError:
                    pass
        for c in candidates:
            if os.path.isfile(c):
                return c
    if sys.platform == "darwin":
        for p in ("/Applications/LibreOffice.app/Contents/MacOS/soffice",):
            if os.path.isfile(p):
                return p
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    return None


def _lo_profile_url(profile_dir: Path) -> str:
    """Lag LibreOffice UserInstallation-URL for isolert profil."""
    return profile_dir.as_uri()  # gir riktig file:///C:/path/... på alle plattformer


def _lo_convert_batch(
    files: list[Path],
    dst_dir: Path,
    to_format: str,
    lo_exe: str,
    timeout_per_file: int,
    profile_base: Path,
    w,
) -> dict[Path, Path]:
    """
    Konverter flere filer i ett enkelt LO-kall (én oppstart for N filer).
    Returnerer {input_path: output_path} for vellykket konverterte filer.
    """
    if not files:
        return {}
    dst_dir.mkdir(parents=True, exist_ok=True)
    profile_dir = Path(tempfile.mkdtemp(dir=profile_base, prefix="lo_prof_"))
    ext    = to_format.split(":")[0]
    result: dict[Path, Path] = {}
    try:
        cmd = [
            lo_exe,
            f"-env:UserInstallation={_lo_profile_url(profile_dir)}",
            "--headless", "--norestore",
            "--convert-to", to_format,
            "--outdir", str(dst_dir),
        ] + [str(f) for f in files]
        subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout_per_file * max(len(files), 1),
        )
        for f in files:
            out = dst_dir / f"{f.stem}.{ext}"
            if not out.exists():
                for c in dst_dir.iterdir():
                    if (c.suffix.lower() == f".{ext}"
                            and c.stem.lower() == f.stem.lower()):
                        out = c
                        break
            if out.exists():
                result[f] = out
    except subprocess.TimeoutExpired:
        w(f"    LO batch-timeout "
          f"({timeout_per_file * len(files)}s, {len(files)} filer)", "warn")
    except Exception as exc:
        w(f"    LO batch-feil: {exc}", "warn")
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)
    return result


def _lo_convert(
    src: Path,
    dst_dir: Path,
    to_format: str,       # f.eks. "docx" eller "docx:MS Word 2007 XML"
    lo_exe: str,
    timeout: int,
    profile_base: Path,   # mappe under hvilken en unik per-kall profil opprettes
    w,
) -> Path | None:
    """
    Konverter én fil med LibreOffice headless.
    Hver kall får sin egen isolerte profil-mappe for å unngå låsekonflikter
    når LO kjøres sekvensielt (bootstrap.ini-feil).
    Returnerer Path til output-filen, eller None ved feil.
    """
    dst_dir.mkdir(parents=True, exist_ok=True)
    # Unik profil per kall — LO låser profilen og frigir den ikke alltid rent
    profile_dir = Path(tempfile.mkdtemp(dir=profile_base, prefix="lo_prof_"))
    try:
        cmd = [
            lo_exe,
            f"-env:UserInstallation={_lo_profile_url(profile_dir)}",
            "--headless", "--norestore",
            "--convert-to", to_format,
            "--outdir", str(dst_dir),
            str(src),
        ]
        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout,
        )
        # LibreOffice skriver output-filen med samme stamme som input
        stem = src.stem
        # format kan være "docx:MS Word 2007 XML" — ta bare første del
        ext  = to_format.split(":")[0]
        out  = dst_dir / f"{stem}.{ext}"
        if out.exists():
            return out
        # Sjekk hva som faktisk ble skrevet (LO kan endre filnavn)
        for f in dst_dir.iterdir():
            if f.suffix.lower() == f".{ext}" and f.stem.lower() == stem.lower():
                return f
        w(f"    LO convert: {src.name} → {to_format}: ingen output-fil funnet. "
          f"stderr: {result.stderr.decode(errors='replace')[:200]}", "warn")
        return None
    except subprocess.TimeoutExpired:
        w(f"    LO convert timeout ({timeout}s): {src.name}", "warn")
        return None
    except Exception as exc:
        w(f"    LO convert feil ({src.name}): {exc}", "warn")
        return None
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)


# ── Avkryptering ──────────────────────────────────────────────────────────────

def _decrypt_file(src: Path, dst: Path, password: str, w) -> bool:
    """
    Avkrypter passordbeskyttet Office-fil med msoffcrypto-tool.
    Returnerer True hvis vellykket, False ellers.

    Faller tilbake til å kopiere filen uendret hvis filen ikke er kryptert
    (msoffcrypto gir OfficeFileError i det tilfellet).
    """
    try:
        import msoffcrypto  # type: ignore
        with open(src, "rb") as f_in:
            office_file = msoffcrypto.OfficeFile(f_in)
            if not office_file.is_encrypted():
                # Filen er ikke kryptert — kopier som-er
                shutil.copy2(src, dst)
                w(f"    {src.name}: ikke kryptert — kopierer uendret", "info")
                return True
            office_file.load_key(password=password)
            with open(dst, "wb") as f_out:
                office_file.decrypt(f_out)
        return True
    except ImportError:
        w("    msoffcrypto-tool ikke installert — kopier fil uavkryptert. "
          "Installer med: pip install msoffcrypto-tool", "warn")
        shutil.copy2(src, dst)
        return False
    except Exception as exc:
        w(f"    Avkryptering feilet ({src.name}): {exc}", "warn")
        # Kopier uendret slik at filen i det minste er med i output
        shutil.copy2(src, dst)
        return False


# ── Mailmerge ─────────────────────────────────────────────────────────────────

def _read_merge_data_from_docx(docx_path: Path, w) -> dict[str, str]:
    """
    Les merge-data fra første tabell i DOCX-filen.
    Rad 0 = feltnavn (kolonneoverskrifter), rad 1 = dataverdier.
    Returnerer dict {feltnavn: verdi} eller tom dict ved feil.
    """
    try:
        import docx  # type: ignore  (python-docx)
        doc = docx.Document(str(docx_path))
        if not doc.tables:
            w(f"    {docx_path.name}: ingen tabeller funnet i datakilde", "warn")
            return {}
        table = doc.tables[0]
        if len(table.rows) < 2:
            w(f"    {docx_path.name}: datakilde-tabell har færre enn 2 rader", "warn")
            return {}
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        values  = [cell.text.strip() for cell in table.rows[1].cells]
        data = {h: v for h, v in zip(headers, values) if h}
        w(f"    Merge-data: {len(data)} felt: {', '.join(list(data.keys())[:6])}"
          f"{'…' if len(data) > 6 else ''}", "info")
        return data
    except ImportError:
        w("    python-docx ikke installert — kan ikke lese merge-data", "warn")
        return {}
    except Exception as exc:
        w(f"    Lesing av merge-data feilet: {exc}", "warn")
        return {}


def _read_merge_data_via_html(
    src_doc: Path,
    lo_exe: str,
    lo_timeout: int,
    profile_base: Path,
    w,
) -> dict[str, str]:
    """
    HTML-fallback: konverter datakilde til HTML med LibreOffice og les tabellen.
    Prøver først å lese <table>-element; faller tilbake til semikolon-separert
    tekst (CosDoc-format: to avsnitt med semikolon-delte feltnavn / verdier).
    """
    try:
        html_dir = src_doc.parent / "html_out"
        html_dir.mkdir(exist_ok=True)
        html_file = _lo_convert(src_doc, html_dir, "html",
                                lo_exe, lo_timeout, profile_base, w)
        if not html_file or not html_file.exists():
            w("    HTML-konvertering ga ingen fil", "warn")
            return {}
        content = html_file.read_text(encoding="utf-8", errors="replace")

        # ── Forsøk 1: HTML-tabell ─────────────────────────────────────────────
        table_m = re.search(r'<table\b[^>]*>(.*?)</table>',
                            content, re.DOTALL | re.IGNORECASE)
        if table_m:
            def _cells(row_html: str) -> list[str]:
                return [
                    re.sub(r'<[^>]+>', '', c).strip()
                    for c in re.findall(
                        r'<t[dh]\b[^>]*>(.*?)</t[dh]>',
                        row_html, re.DOTALL | re.IGNORECASE)
                ]
            rows = re.findall(r'<tr\b[^>]*>(.*?)</tr>',
                              table_m.group(1), re.DOTALL | re.IGNORECASE)
            if len(rows) >= 2:
                data = {h: v
                        for h, v in zip(_cells(rows[0]), _cells(rows[1])) if h}
                if data:
                    w(f"    Merge-data (HTML-tabell): {len(data)} felt", "info")
                    return data

        # ── Forsøk 2: semikolon-separert tekst (CosDoc-format) ───────────────
        plain = re.sub(r'<[^>]+>', '\n', content)

        def _is_identifier_line(line: str) -> bool:
            """Sjekk at alle semikolon-delte deler ser ut som feltidentifikatorer."""
            parts = [p.strip() for p in line.split(';') if p.strip()]
            return bool(parts) and all(re.match(r'^[A-Za-z_]\w*$', p) for p in parts)

        lines = [ln.strip() for ln in plain.splitlines()
                 if ';' in ln and ln.strip() and _is_identifier_line(ln)]
        if len(lines) >= 1:
            # Finn tilsvarende datalinje (første linje som ikke er identifikatorer)
            plain_lines = [ln.strip() for ln in plain.splitlines()
                           if ';' in ln and ln.strip()]
            header_line = lines[0]
            header_idx  = plain_lines.index(header_line) if header_line in plain_lines else -1
            value_line  = plain_lines[header_idx + 1] if header_idx >= 0 and header_idx + 1 < len(plain_lines) else ""
            if value_line:
                headers = [h.strip() for h in header_line.split(';')]
                values  = [v.strip() for v in value_line.split(';')]
                data    = {h: v for h, v in zip(headers, values) if h}
                if data:
                    w(f"    Merge-data (HTML-tekst): {len(data)} felt: "
                      f"{', '.join(list(data.keys())[:6])}"
                      f"{'…' if len(data) > 6 else ''}", "info")
                    return data

        w("    Ingen tabell eller semikolondata i HTML-output", "warn")
        return {}
    except Exception as exc:
        w(f"    HTML-fallback feilet: {exc}", "warn")
        return {}


def _read_merge_data_from_doc_binary(doc_path: Path, w) -> dict[str, str]:
    """
    Les merge-data direkte fra Word 97-2003 OLE2-binærformat.

    CosDoc-datatkildefiler (.doc) lagrer data som semikolon-separert tekst:
      avsnitt 0 = feltnavn  (eks. "Sdo_DokID;Sdo_ArkivSakID;...")
      avsnitt 1 = verdier   (eks. "2013001733;2013000105;...")

    Bruker FIB-feltene fcMin/fcMac til å lokalisere teksten i WordDocument-
    strømmen uten ekstern parsing-avhengighet (olefile følger med msoffcrypto).
    """
    try:
        import olefile as _ole  # type: ignore  (følger med msoffcrypto-tool)
        if not _ole.isOleFile(str(doc_path)):
            return {}
        ole = _ole.OleFileIO(str(doc_path))
        try:
            if not ole.exists('WordDocument'):
                return {}
            wd = ole.openstream('WordDocument').read()
        finally:
            ole.close()

        # FIB magic = 0xA5EC
        if int.from_bytes(wd[0:2], 'little') != 0xA5EC:
            return {}

        fcMin = int.from_bytes(wd[24:28], 'little')
        fcMac = int.from_bytes(wd[28:32], 'little')
        if not (0 <= fcMin < fcMac <= len(wd)):
            return {}

        text_bytes = wd[fcMin:fcMac]
        # Prøv cp1252 (Windows ANSI, norm for norske .doc-filer), deretter UTF-16LE
        try:
            text = text_bytes.decode('cp1252', errors='strict')
        except (UnicodeDecodeError, ValueError):
            text = text_bytes.decode('utf-16-le', errors='replace')

        # Avsnitt separert med CR (0x0D); tabellceller med 0x07
        paras = [p.strip() for p in re.split(r'[\x07\x0d]', text) if p.strip()]
        if len(paras) < 2:
            return {}

        # Finn første avsnitt med semikolon (feltnavn-rad)
        header_idx = next((i for i, p in enumerate(paras) if ';' in p), None)
        if header_idx is None or header_idx + 1 >= len(paras):
            return {}

        headers = [h.strip() for h in paras[header_idx].split(';')]
        values  = [v.strip() for v in paras[header_idx + 1].split(';')]
        data    = {h: v for h, v in zip(headers, values) if h}
        if data:
            w(f"    Merge-data (DOC-binær): {len(data)} felt: "
              f"{', '.join(list(data.keys())[:6])}"
              f"{'…' if len(data) > 6 else ''}", "info")
        return data
    except ImportError:
        return {}
    except Exception as exc:
        w(f"    DOC-binær lesing feilet: {exc}", "warn")
        return {}


def _perform_mailmerge(
    main_docx: Path,
    merge_data: dict[str, str],
    output: Path,
    w,
) -> bool:
    """
    Utfør mailmerge på main_docx med data fra merge_data.
    Skriver resultat til output.

    Bruker docx-mailmerge2 (pip install docx-mailmerge2).
    Returnerer True hvis vellykket, False ellers.
    """
    try:
        from mailmerge import MailMerge  # type: ignore  (docx-mailmerge2)
        with MailMerge(str(main_docx)) as document:
            fields = document.get_merge_fields()
            if not fields:
                w(f"    {main_docx.name}: ingen MERGEFIELD-felt funnet", "warn")
                return False
            # Filtrer data til kun felter som finnes i dokumentet
            filtered = {k: v for k, v in merge_data.items() if k in fields}
            missing  = fields - set(filtered.keys())
            if missing:
                w(f"    MERGEFIELD-felt uten data: {', '.join(sorted(missing))}", "warn")
            document.merge(**filtered)
            document.write(str(output))
        w(f"    Flettet {len(filtered)} felt → {output.name}", "ok")
        return True
    except ImportError:
        w("    docx-mailmerge2 ikke installert — hopper over fletting. "
          "Installer med: pip install docx-mailmerge2", "warn")
        return False
    except Exception as exc:
        w(f"    Mailmerge feilet: {exc}", "warn")
        return False


# ── DOCX datakilde-patching ───────────────────────────────────────────────────

def _patch_docx_datasource(docx_path: Path, new_src: Path) -> None:
    """
    Oppdater datakilde-stien i word/settings.xml i en DOCX så fletteverktøy
    kan finne flettefilen i temp-mappen (unngår feil på hardkodet produksjonssti).

    Håndterer to vanlige OOXML-formater:
      <w:dataSource w:val="..."/>
      <w:connectString w:val="...DBQ=path\to\file..."/>
    """
    SETTINGS = "word/settings.xml"
    try:
        buf = io.BytesIO()
        new_uri  = new_src.as_uri()          # file:///C:/...
        new_path = str(new_src)              # C:\...
        changed  = False

        with zipfile.ZipFile(docx_path, "r") as zin, \
             zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED, allowZip64=True) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == SETTINGS:
                    orig = data
                    # Format 1: <w:dataSource w:val="file:///..."/>
                    data = re.sub(
                        rb'(w:dataSource\s[^>]*w:val=")[^"]*(")',
                        lambda m: m.group(1) + new_uri.encode() + m.group(2),
                        data,
                    )
                    # Format 2: <w:connectString w:val="...DBQ=path\file..."/>
                    data = re.sub(
                        rb'(DBQ=)[^;"\s<]+',
                        lambda m: m.group(1) + new_path.encode(),
                        data,
                    )
                    changed = data != orig
                zout.writestr(item, data)
        if changed:
            docx_path.write_bytes(buf.getvalue())
    except Exception:
        pass  # ikke kritisk — fletting forsøkes uansett


# ── SIARD ZIP-hjelper ─────────────────────────────────────────────────────────

def _orig_dir_entries(namelist: list[str]) -> list[str]:
    return sorted(n for n in namelist if n.endswith("/"))


def _pack_zip(
    extract_dir: Path,
    orig_namelist: list[str],
    dst_path: Path,
    target_version: str,
    w,
) -> None:
    """Pakk om extract_dir til ny SIARD ZIP med korrekte katalogoppføringer."""
    dir_entries = _orig_dir_entries(orig_namelist)
    n_written = n_transformed = 0

    with zipfile.ZipFile(dst_path, "w", zipfile.ZIP_DEFLATED, allowZip64=True) as zf:
        for de in dir_entries:
            zf.writestr(zipfile.ZipInfo(de), b"")
            n_written += 1

        for fp in sorted(extract_dir.rglob("*")):
            if not fp.is_file():
                continue
            arc = str(fp.relative_to(extract_dir)).replace("\\", "/")
            if is_siard_xml(arc):
                data = fp.read_bytes()
                data = siard_version_transform(data, target_version)
                zf.writestr(arc, data)
                n_transformed += 1
            else:
                zf.write(fp, arc)
            n_written += 1

    if n_transformed:
        w(f"  {n_transformed} XML-filer transformert til SIARD {target_version}",
          "info")
    w(f"  Pakket: {n_written:,} oppføringer  ({dst_path.stat().st_size:,} bytes)",
      "ok")


# ── Oppdatering av tableX.xml ─────────────────────────────────────────────────

def _update_table_xml(
    xml_bytes: bytes,
    updates: dict[int, dict],    # {eef_id: {file, length, digest, new_filename}}
    deletes: set[int],           # {eef_id} for rader som skal slettes
    eefid_col: int,
    filename_col: int,
    blob_col: int,
    size_col: int,
    w,
) -> bytes:
    """
    Oppdater tableX.xml:
      - updates: oppdater blob-referanse, filnavn og størrelse for hoveddokumenter
      - deletes: fjern datakilde-rader fra XML

    Returnerer oppdaterte XML-bytes.
    Bruker byte-nivå serialisering via ET for å bevare XML-deklarasjon.
    """
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError as exc:
        w(f"    XML-parse-feil: {exc}", "feil")
        return xml_bytes

    rows_to_remove = []
    for row_el in list(root):
        if _local(row_el.tag) != "row":
            continue
        # Finn Eef_EefID for denne raden
        eef_id_str = ""
        for child in row_el:
            tag = _local(child.tag)
            m = re.match(r"^c(\d+)$", tag, re.IGNORECASE)
            if m and int(m.group(1)) == eefid_col:
                eef_id_str = (child.text or "").strip()
                break

        try:
            eef_id = int(eef_id_str)
        except ValueError:
            continue

        if eef_id in deletes:
            rows_to_remove.append(row_el)
            continue

        if eef_id in updates:
            upd = updates[eef_id]
            for child in row_el:
                tag = _local(child.tag)
                m = re.match(r"^c(\d+)$", tag, re.IGNORECASE)
                if not m:
                    continue
                idx = int(m.group(1))
                if idx == blob_col:
                    child.set("file",       upd["file"])
                    child.set("length",     str(upd["length"]))
                    child.set("digestType", "MD5")
                    child.set("digest",     upd["digest"])
                    child.text = None
                elif idx == filename_col:
                    child.text = upd["new_filename"]
                elif idx == size_col:
                    child.text = str(upd["length"])

    for row_el in rows_to_remove:
        root.remove(row_el)

    w(f"    XML: {len(updates)} rad(er) oppdatert, "
      f"{len(rows_to_remove)} rad(er) fjernet", "info")

    # Serialiser tilbake
    ET.indent(root, space="\t")
    tree = ET.ElementTree(root)
    buf = io.BytesIO()
    tree.write(buf, encoding="utf-8", xml_declaration=True)
    return buf.getvalue()


# ── Hoved-operasjon ───────────────────────────────────────────────────────────

class CosDocMailMergeOperation(BaseOperation):
    """
    System-spesifikk operasjon for CosDoc-fagsystem.

    Avkrypterer passordbeskyttede dokumenter fra Eef_ElFiler-tabellen og
    utfører mailmerge for dokumentpar (Eef_SeqNr 1 + 2 med samme Eef_EveID).

    Krav til kjøretidsmiljø:
      - LibreOffice installert (for .doc → .docx konvertering)
      - msoffcrypto-tool  (pip install msoffcrypto-tool)  — for avkryptering
      - docx-mailmerge2   (pip install docx-mailmerge2)   — for mailmerge
      - python-docx       (installeres av docx-mailmerge2) — for datalesing
    """

    operation_id   = "cosdoc_mailmerge"
    label          = "CosDoc: Lås opp og flett dokumenter"
    description    = (
        "CosDoc-spesifikk: Avkrypterer passordbeskyttede dokumenter i "
        "Eef_ElFiler-tabellen og utfører mailmerge for dokumentpar "
        "(Eef_SeqNr 1+2 med samme Eef_EveID). "
        "Passord utledes av filnavnet (R + omvendt filstamme). "
        "Krever msoffcrypto-tool og docx-mailmerge2."
    )
    category       = "Systemspesifikt"
    status         = 2
    produces_siard = True

    default_params: dict = {
        "dry_run":        False,
        "output_suffix":  "_cosdoc",
        "table_name":     "Eef_ElFiler",
        "lo_executable":  "",      # soffice-sti — auto-detekteres hvis tom
        "lo_timeout":     120,     # sekunder per LO-konvertering
        "temp_dir":       "",      # temp-rotmappe
    }

    def run(self, ctx: WorkflowContext) -> OperationResult:  # noqa: C901
        log = ctx.metadata.get("file_logger")
        pcb = ctx.metadata.get("progress_cb")

        def w(msg: str, lvl: str = "info") -> None:
            if log:
                log.log(msg, lvl)
            if lvl in ("ok", "warn", "feil", "step", "info"):
                if pcb:
                    pcb("log", msg=msg, level=lvl)

        def progress(event: str, **kw) -> None:
            if pcb:
                pcb(event, **kw)

        # ── Pipeline-modus: UnpackSiardOperation har allerede pakket ut ────────
        pipeline_mode = bool(
            getattr(ctx, "extracted_path", None)
            and ctx.extracted_path.is_dir()
        )
        if pipeline_mode:
            self.produces_siard = False
            PHASES = 3
            w("  Pipeline-modus: jobber direkte på utpakket filstruktur "
              "(repakking via 'Pakk sammen SIARD')", "info")
        else:
            PHASES = 5

        def phase(n: int, label: str) -> None:
            progress("phase", phase=n, total_phases=PHASES, label=label)

        # ── Parametere ────────────────────────────────────────────────────────
        dry_run    = bool(self.params["dry_run"])
        suffix     = str(self.params["output_suffix"] or "_cosdoc")
        table_name = str(self.params["table_name"] or "Eef_ElFiler")
        from settings import get_config as _get_cfg
        lo_hint       = str(self.params["lo_executable"] or _get_cfg("lo_executable", "") or "soffice")
        lo_timeout    = int(self.params["lo_timeout"] or 120)
        max_workers   = max(1, min(int(_get_cfg("max_workers",  4)  or 4), os.cpu_count() or 4))
        lo_batch_size = max(1, int(_get_cfg("lo_batch_size", 20) or 20))

        temp_root_str = str(self.params.get("temp_dir", "") or "")
        if not temp_root_str and hasattr(ctx, "metadata"):
            temp_root_str = ctx.metadata.get("temp_dir", "")
        temp_root = Path(temp_root_str) if temp_root_str else None

        src_path = ctx.siard_path
        if src_path is None:
            return self._fail("Ingen SIARD-fil angitt i context.")
        if not pipeline_mode and not src_path.exists():
            return self._fail(f"SIARD-fil ikke funnet: {src_path}")

        dst_path = src_path.with_stem(src_path.stem + suffix)
        output_dir_override = ""
        if hasattr(ctx, "metadata"):
            output_dir_override = ctx.metadata.get("output_dir_override", "").strip()
        if output_dir_override:
            dst_path = Path(output_dir_override) / dst_path.name

        # ── Fase 1: LibreOffice ───────────────────────────────────────────────
        phase(1, "Finner LibreOffice")
        lo_exe = _find_libreoffice(lo_hint)
        if not lo_exe:
            return self._fail(
                "LibreOffice ikke funnet. Installer fra https://www.libreoffice.org "
                "og kontroller at soffice er i PATH.")
        w(f"  LibreOffice: {lo_exe}", "info")
        progress("phase_done")

        stats = {
            "pairs_found":       0,
            "pairs_merged":      0,
            "singles_found":     0,
            "singles_decrypted": 0,
            "decrypt_failed":    0,
            "merge_failed":      0,
        }

        with tempfile.TemporaryDirectory(dir=temp_root) as _tmpdir:
            tmpdir     = Path(_tmpdir)
            lo_profile_base = tmpdir / "lo_profiles"  # base for per-kall profiler
            work_dir        = tmpdir / "work"
            lo_profile_base.mkdir()
            work_dir.mkdir()

            # ── Fase 2 (standalone): Pakk ut SIARD ───────────────────────────
            if pipeline_mode:
                extract_dir    = ctx.extracted_path
                orig_namelist  = []   # ikke brukt i pipeline-modus
                target_version = get_target_siard_version()
            else:
                phase(2, "Pakker ut SIARD-arkiv")
                extract_dir = tmpdir / "extracted"
                extract_dir.mkdir()
                try:
                    with zipfile.ZipFile(src_path, "r") as zf:
                        orig_namelist = zf.namelist()
                        meta_arc = next(
                            (n for n in orig_namelist
                             if n.lower().endswith("header/metadata.xml")), None)
                        if not meta_arc:
                            return self._fail(
                                "metadata.xml ikke funnet i SIARD-arkivet.")
                        src_version    = detect_siard_version(zf.read(meta_arc))
                        target_version = get_target_siard_version()
                        w(f"  SIARD: kilde={src_version}, mål={target_version}",
                          "info")
                        for name in orig_namelist:
                            try:
                                zf.extract(name, extract_dir)
                            except Exception as exc:
                                w(f"  [ADVARSEL] Utpakking {name}: {exc}", "warn")
                except Exception as exc:
                    return self._fail(f"Kan ikke åpne SIARD-arkiv: {exc}")
                w(f"  Pakket ut {len(orig_namelist)} oppføringer", "info")
                progress("phase_done")

            # ── Fase 2 (pipeline) / Fase 3 (standalone): Analyser tabell ─────
            phase(2 if pipeline_mode else 3, f"Analyserer {table_name}")

            metadata_path = extract_dir / "header" / "metadata.xml"
            if not metadata_path.exists():
                return self._fail(
                    "header/metadata.xml ikke funnet i utpakket arkiv.")

            table_info = _find_table_info(metadata_path.read_bytes(), table_name)
            if not table_info:
                return self._fail(
                    f"Tabell '{table_name}' ikke funnet i metadata.xml. "
                    "Kontroller tabell-navn i operasjonsinnstillingene.")

            sf   = table_info["schema_folder"]
            tf   = table_info["table_folder"]
            cm   = table_info["col_map"]
            lobs = table_info["lob_folders"]

            w(f"  Tabell: {sf}/{tf}  ({len(cm)} kolonner, "
              f"{len(lobs)} BLOB-kolonner)", "info")

            required = ("Eef_EefID", "Eef_EveID", "Eef_SeqNr",
                        "Eef_FilNavn", "Eef_Size", "Eef_ElFil")
            missing_cols = [c for c in required if c not in cm]
            if missing_cols:
                return self._fail(
                    f"Manglende kolonner i {table_name}: "
                    + ", ".join(missing_cols))

            eefid_col  = cm["Eef_EefID"]
            eveid_col  = cm["Eef_EveID"]
            seqnr_col  = cm["Eef_SeqNr"]
            fname_col  = cm["Eef_FilNavn"]
            size_col   = cm["Eef_Size"]
            blob_col   = cm["Eef_ElFil"]
            lob_folder = lobs.get(blob_col, f"{sf}/{tf}/lob{blob_col}")

            table_xml_path = extract_dir / "content" / sf / tf / f"{tf}.xml"
            if not table_xml_path.exists():
                return self._fail(f"tableX.xml ikke funnet: {table_xml_path}")

            xml_bytes = table_xml_path.read_bytes()
            rows      = _parse_table_rows(xml_bytes, cm)
            w(f"  Rader i {tf}.xml: {len(rows)}", "info")

            by_eveid: dict[str, list[dict]] = {}
            for row in rows:
                eve_id = _get_col(row, eveid_col)
                if eve_id:
                    by_eveid.setdefault(eve_id, []).append(row)

            pairs:   list[tuple[dict, dict]] = []
            singles: list[dict]              = []
            for eve_rows in by_eveid.values():
                main_rows = [r for r in eve_rows if _get_col(r, seqnr_col) == "1"]
                data_rows = [r for r in eve_rows if _get_col(r, seqnr_col) == "2"]
                for main_row in main_rows:
                    if data_rows:
                        pairs.append((main_row, data_rows[0]))
                    else:
                        singles.append(main_row)

            stats["pairs_found"]   = len(pairs)
            stats["singles_found"] = len(singles)
            w(f"  Dokumentpar (SeqNr 1+2): {len(pairs)}, "
              f"enkeltdokumenter (SeqNr 1): {len(singles)}", "info")
            progress("phase_done")

            # ── Fase 3 / 4: Avkrypterer, konverterer, fletter ────────────────
            phase(3 if pipeline_mode else 4, "Avkrypterer og fletter dokumenter")

            _lf = Path(lob_folder)
            if _lf.parts and _lf.parts[0].lower() == "content":
                lob_path = extract_dir / _lf
            else:
                lob_path = extract_dir / "content" / _lf

            xml_updates: dict[int, dict] = {}
            xml_deletes: set[int]        = set()
            total        = len(pairs) + len(singles)
            lock         = threading.Lock()
            done_count   = [0]
            lo_out_root   = work_dir / "lo_out"
            lo_out_root.mkdir(exist_ok=True)

            w(f"  Parallellitet: {max_workers} tråder, "
              f"LO batch-størrelse: {lo_batch_size}", "info")

            # ── Steg A: Avkrypter + les binær merge-data (parallell, ingen LO) ─

            class _PP:
                """Forberedt par-objekt."""
                __slots__ = ("seq", "eef_id_main", "eef_id_data", "eveid",
                             "main_fname", "data_fname", "main_blob", "data_blob",
                             "main_dec", "data_dec", "merge_data",
                             "needs_data_lo", "skip", "sd")

            class _SP:
                """Forberedt enkelt-objekt."""
                __slots__ = ("seq", "eef_id", "fname", "blob",
                             "dec_path", "is_word", "skip", "sd")

            def _prep_pair(seq: int, main_row: dict, data_row: dict) -> _PP:
                p               = _PP()
                p.seq           = seq
                p.eef_id_main   = int(_get_col(main_row, eefid_col) or "0")
                p.eef_id_data   = int(_get_col(data_row,  eefid_col) or "0")
                p.eveid         = _get_col(main_row, eveid_col)
                p.main_fname    = _get_col(main_row, fname_col)
                p.data_fname    = _get_col(data_row,  fname_col)
                p.main_blob     = _get_blob(main_row, blob_col)
                p.data_blob     = _get_blob(data_row,  blob_col)
                p.merge_data    = {}
                p.needs_data_lo = False
                p.main_dec      = None
                p.data_dec      = None
                p.skip          = False
                p.sd            = {}

                w(f"\n  [{seq}/{total}] Par EveID={p.eveid}: "
                  f"{p.main_fname} + {p.data_fname}", "step")

                if not p.main_blob or not p.data_blob:
                    w("    HOPPER OVER: manglende blob-referanse", "warn")
                    p.skip = True
                    return p
                main_src = lob_path / p.main_blob["file"]
                data_src = lob_path / p.data_blob["file"]
                if not main_src.exists():
                    w(f"    FEIL: blob ikke funnet: {main_src}", "feil")
                    p.skip = True
                    return p
                if dry_run:
                    p.sd["pairs_merged"] = 1
                    p.skip = True
                    return p

                pair_dir   = work_dir / f"pair_{p.eef_id_main}"
                pair_dir.mkdir(exist_ok=True)
                p.main_dec = pair_dir / p.main_fname
                p.data_dec = pair_dir / p.data_fname

                main_pw = _derive_password(p.main_fname)
                w(f"    Avkrypterer {p.main_fname} (pw={main_pw[:4]}…)", "info")
                if not _decrypt_file(main_src, p.main_dec, main_pw, w):
                    p.sd["decrypt_failed"] = 1

                if data_src.exists():
                    shutil.copy2(data_src, p.data_dec)
                    w(f"    Kopierer datakilde {p.data_fname}", "info")
                else:
                    w(f"    Datakilde ikke funnet: {data_src}", "warn")

                if p.data_dec.exists():
                    p.merge_data = _read_merge_data_from_doc_binary(p.data_dec, w)
                    if not p.merge_data:
                        p.needs_data_lo = True
                return p

            def _prep_single(seq: int, row: dict) -> _SP:
                s          = _SP()
                s.seq      = seq
                s.eef_id   = int(_get_col(row, eefid_col) or "0")
                s.fname    = _get_col(row, fname_col)
                s.blob     = _get_blob(row, blob_col)
                s.is_word  = _is_word_doc(s.fname)
                s.dec_path = None
                s.skip     = False
                s.sd       = {}

                w(f"\n  [{seq}/{total}] Enkelt: {s.fname}", "step")

                if not s.blob:
                    w("    HOPPER OVER: manglende blob-referanse", "warn")
                    s.skip = True
                    return s
                blob_src = lob_path / s.blob["file"]
                if not blob_src.exists():
                    w(f"    FEIL: blob ikke funnet: {blob_src}", "feil")
                    s.skip = True
                    return s
                if dry_run:
                    s.sd["singles_decrypted"] = 1
                    s.skip = True
                    return s

                single_dir = work_dir / f"single_{s.eef_id}"
                single_dir.mkdir(exist_ok=True)
                s.dec_path = single_dir / s.fname

                if s.is_word:
                    pw = _derive_password(s.fname)
                    w(f"    Avkrypterer {s.fname} (pw={pw[:4]}…)", "info")
                    if not _decrypt_file(blob_src, s.dec_path, pw, w):
                        s.sd["decrypt_failed"] = 1
                else:
                    shutil.copy2(blob_src, s.dec_path)
                return s

            pair_preps:   list[_PP] = []
            single_preps: list[_SP] = []
            with concurrent.futures.ThreadPoolExecutor(
                    max_workers=max_workers) as pool:
                pfuts = {pool.submit(_prep_pair,   i + 1,             mr, dr): i
                         for i, (mr, dr) in enumerate(pairs)}
                sfuts = {pool.submit(_prep_single, len(pairs) + i + 1, r):  i
                         for i, r in enumerate(singles)}
                all_futs = {**pfuts, **sfuts}
                for fut in concurrent.futures.as_completed(all_futs):
                    try:
                        res = fut.result()
                    except Exception as exc:
                        w(f"  Prep-feil: {exc}", "feil")
                        continue
                    with lock:
                        for k, v in res.sd.items():
                            stats[k] = stats.get(k, 0) + v
                    if fut in pfuts:
                        pair_preps.append(res)
                    else:
                        single_preps.append(res)

            # ── Steg B: Batch LO-konvertering ─────────────────────────────────
            # Samle alle filer som trenger konvertering
            to_convert: list[Path] = []
            for p in pair_preps:
                if (not p.skip and p.main_dec and p.main_dec.exists()
                        and p.main_dec.suffix.lower() != ".docx"):
                    to_convert.append(p.main_dec)
            for p in pair_preps:
                if (not p.skip and p.needs_data_lo
                        and p.data_dec and p.data_dec.exists()
                        and p.data_dec.suffix.lower() != ".docx"):
                    to_convert.append(p.data_dec)
            for s in single_preps:
                if (not s.skip and s.is_word
                        and s.dec_path and s.dec_path.exists()
                        and s.dec_path.suffix.lower() != ".docx"):
                    to_convert.append(s.dec_path)

            converted: dict[Path, Path] = {}  # input_path → output_docx
            if to_convert:
                batches   = [to_convert[s:s + lo_batch_size]
                             for s in range(0, len(to_convert), lo_batch_size)]
                n_batches = len(batches)
                w(f"\n  LO-konvertering: {len(to_convert)} filer, "
                  f"{n_batches} batch(er) à {lo_batch_size}, "
                  f"{max_workers} parallelle LO-prosesser", "step")

                def _run_batch(args: tuple) -> dict[Path, Path]:
                    bi, batch = args
                    batch_dir = lo_out_root / f"b{bi}"
                    batch_dir.mkdir(exist_ok=True)
                    res = _lo_convert_batch(
                        batch, batch_dir, "docx",
                        lo_exe, lo_timeout, lo_profile_base, w)
                    w(f"  Batch {bi + 1}/{n_batches}: "
                      f"{len(res)}/{len(batch)} OK", "info")
                    return res

                with concurrent.futures.ThreadPoolExecutor(
                        max_workers=max_workers) as pool:
                    for batch_res in pool.map(_run_batch, enumerate(batches)):
                        converted.update(batch_res)

            # ── Steg C: Flett og skriv (parallell, ingen LO-kall) ─────────────

            def _merge_pair(p: _PP) -> dict:
                base = {"type": "pair", "eef_id_main": p.eef_id_main,
                        "eef_id_data": p.eef_id_data,
                        "update": None, "delete": False, "sd": dict(p.sd)}
                if p.skip:
                    return base

                main_docx = converted.get(p.main_dec)

                if not p.merge_data and p.needs_data_lo:
                    data_docx = converted.get(p.data_dec) if p.data_dec else None
                    if data_docx:
                        p.merge_data = _read_merge_data_from_docx(data_docx, w)
                    if not p.merge_data and p.data_dec and p.data_dec.exists():
                        p.merge_data = _read_merge_data_via_html(
                            p.data_dec, lo_exe, lo_timeout, lo_profile_base, w)

                if main_docx and p.data_dec and p.data_dec.exists():
                    _patch_docx_datasource(main_docx, p.data_dec)

                merged_path = p.main_dec.parent / f"merged_{p.eef_id_main}.docx"
                merge_ok = False
                if main_docx and p.merge_data:
                    merge_ok = _perform_mailmerge(
                        main_docx, p.merge_data, merged_path, w)
                elif main_docx and not p.merge_data:
                    w("    Ingen merge-data — lagrer konvertert hoveddokument",
                      "warn")

                if not merge_ok:
                    if main_docx and main_docx.exists():
                        shutil.copy2(main_docx, merged_path)
                        w("    Bruker konvertert hoveddokument (uten merge)", "info")
                    elif p.main_dec.exists():
                        merged_path = p.main_dec
                        w("    Bruker avkryptert hoveddokument (uten konvertering)",
                          "info")
                    else:
                        w(f"    Kan ikke produsere output for {p.main_fname}", "feil")
                        base["sd"]["merge_failed"] = 1
                        return base
                    base["sd"]["merge_failed"] = base["sd"].get("merge_failed", 0) + 1
                else:
                    base["sd"]["pairs_merged"] = base["sd"].get("pairs_merged", 0) + 1

                new_fname     = Path(p.main_fname).stem + "_flettet.docx"
                new_blobname  = Path(p.main_blob["file"]).stem + merged_path.suffix
                new_blob_path = lob_path / new_blobname
                merged_bytes  = merged_path.read_bytes()
                new_blob_path.write_bytes(merged_bytes)
                w(f"    Lagret: {new_blobname} ({len(merged_bytes):,} bytes)", "ok")

                old_blob = lob_path / p.main_blob["file"]
                if old_blob != new_blob_path and old_blob.exists():
                    old_blob.unlink()
                data_src_file = lob_path / p.data_blob["file"]
                if data_src_file.exists():
                    data_src_file.unlink()
                    w(f"    Slettet datakilde: {p.data_blob['file']}", "info")

                base["update"] = {"file": new_blobname, "length": len(merged_bytes),
                                  "digest": _md5_upper(merged_bytes),
                                  "new_filename": new_fname}
                base["delete"] = True
                return base

            def _finish_single(s: _SP) -> dict:
                base = {"type": "single", "eef_id": s.eef_id,
                        "update": None, "sd": dict(s.sd)}
                if s.skip:
                    return base

                docx_result = (converted.get(s.dec_path)
                               if (s.is_word and s.dec_path) else None)
                blob_src = lob_path / s.blob["file"]

                if docx_result and docx_result.exists():
                    new_fname     = Path(s.fname).stem + ".docx"
                    new_blobname  = Path(s.blob["file"]).stem + ".docx"
                    new_blob_path = lob_path / new_blobname
                    new_data      = docx_result.read_bytes()
                    new_blob_path.write_bytes(new_data)
                    old_blob = lob_path / s.blob["file"]
                    if old_blob != new_blob_path and old_blob.exists():
                        old_blob.unlink()
                    w(f"    Avkryptert og konvertert: {new_blobname} "
                      f"({len(new_data):,} bytes)", "ok")
                    base["update"] = {"file": new_blobname, "length": len(new_data),
                                      "digest": _md5_upper(new_data),
                                      "new_filename": new_fname}
                    base["sd"]["singles_decrypted"] = 1
                elif s.dec_path and s.dec_path.exists():
                    new_data = s.dec_path.read_bytes()
                    blob_src.write_bytes(new_data)
                    action = "Avkryptert" if s.is_word else "Kopierer"
                    w(f"    {action} (ikke konvertert): {s.blob['file']}", "info")
                    base["update"] = {"file": s.blob["file"], "length": len(new_data),
                                      "digest": _md5_upper(new_data),
                                      "new_filename": s.fname}
                    base["sd"]["singles_decrypted"] = 1
                return base

            with concurrent.futures.ThreadPoolExecutor(
                    max_workers=max_workers) as pool:
                futs = ([pool.submit(_merge_pair,    p) for p in pair_preps]
                        + [pool.submit(_finish_single, s) for s in single_preps])
                for fut in concurrent.futures.as_completed(futs):
                    try:
                        res = fut.result()
                    except Exception as exc:
                        w(f"  Worker-feil: {exc}", "feil")
                        res = {"sd": {}}
                    with lock:
                        for k, v in res.get("sd", {}).items():
                            stats[k] = stats.get(k, 0) + v
                        if res.get("type") == "pair":
                            if res.get("update"):
                                xml_updates[res["eef_id_main"]] = res["update"]
                            if res.get("delete"):
                                xml_deletes.add(res["eef_id_data"])
                        elif res.get("type") == "single" and res.get("update"):
                            xml_updates[res["eef_id"]] = res["update"]
                        done_count[0] += 1
                        progress("phase_progress",
                                 done=done_count[0], total=total)

            if not dry_run and (xml_updates or xml_deletes):
                w(f"\n  Oppdaterer {tf}.xml …", "step")
                updated_xml = _update_table_xml(
                    xml_bytes, xml_updates, xml_deletes,
                    eefid_col, fname_col, blob_col, size_col, w,
                )
                table_xml_path.write_bytes(updated_xml)

            progress("phase_done")

            # ── Fase 5 (standalone): Pakk ny SIARD ───────────────────────────
            if not pipeline_mode:
                phase(5, "Pakker ny SIARD-fil")
                w(f"\n  Destination: {dst_path}", "info")
                if not dry_run:
                    _pack_zip(extract_dir, orig_namelist,
                              dst_path, target_version, w)
                progress("phase_done")

        # ── Oppsummering ──────────────────────────────────────────────────────
        w("", "info")
        w("  ── Oppsummering ─────────────────────────────────────────────",
          "step")
        w(f"  Dokumentpar behandlet: "
          f"{stats['pairs_merged']}/{stats['pairs_found']}",
          "ok" if stats["merge_failed"] == 0 else "warn")
        w(f"  Enkeltdokumenter: "
          f"{stats['singles_decrypted']}/{stats['singles_found']}", "info")
        if stats["decrypt_failed"]:
            w(f"  Avkryptering feilet: {stats['decrypt_failed']}", "warn")
        if stats["merge_failed"]:
            w(f"  Mailmerge feilet/hoppet over: {stats['merge_failed']}", "warn")
        if pipeline_mode:
            w("  Pipeline-modus: repakking overlates til 'Pakk sammen SIARD'",
              "info")

        msg = (
            f"CosDoc prosessert: "
            f"{stats['pairs_merged']}/{stats['pairs_found']} par flettet, "
            f"{stats['singles_decrypted']}/{stats['singles_found']} enkeltdok."
        )
        result_data = {**stats}
        if not pipeline_mode:
            result_data["output_path"] = str(dst_path)
            success = dst_path.exists() or dry_run
        else:
            success = True

        if success:
            return self._ok(result_data, msg)
        return self._fail(f"Ingen output-fil produsert: {dst_path}", result_data)
